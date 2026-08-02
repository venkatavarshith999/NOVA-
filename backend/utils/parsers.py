"""
Multi-format document parsing.
Returns a list of (page_number, text) tuples so downstream chunking can
retain page-level citations.
"""
import io
import os
import csv
from typing import List, Tuple

import fitz  # PyMuPDF
import docx
import openpyxl
import pandas as pd
from PIL import Image


def parse_pdf(path: str) -> List[Tuple[int, str]]:
    pages = []
    doc = fitz.open(path)
    for i, page in enumerate(doc):
        text = page.get_text("text")
        pages.append((i + 1, text))
    doc.close()
    return pages


def parse_docx(path: str) -> List[Tuple[int, str]]:
    d = docx.Document(path)
    full_text = []
    for para in d.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                full_text.append(" | ".join(cells))
    return [(1, "\n".join(full_text))]


def parse_txt(path: str) -> List[Tuple[int, str]]:
    with open(path, "r", errors="ignore") as f:
        return [(1, f.read())]


def parse_csv(path: str) -> List[Tuple[int, str]]:
    rows = []
    with open(path, "r", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return [(1, "\n".join(rows))]


def parse_xlsx(path: str) -> List[Tuple[int, str]]:
    pages = []
    wb = openpyxl.load_workbook(path, data_only=True)
    for i, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        lines = [f"Sheet: {sheet_name}"]
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) if v is not None else "" for v in row]
            if any(vals):
                lines.append(" | ".join(vals))
        pages.append((i + 1, "\n".join(lines)))
    return pages


def parse_image(path: str) -> List[Tuple[int, str]]:
    """
    OCR extraction. Uses pytesseract if the tesseract binary is available on
    the host; otherwise records the image as a visual asset with a
    descriptive stub so the pipeline still completes end-to-end.
    """
    try:
        import pytesseract
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        if text.strip():
            return [(1, text)]
    except Exception:
        pass
    return [(1, f"[Image document: {os.path.basename(path)} — no OCR engine available in this environment. "
                 f"Deploy with EasyOCR/Tesseract installed to extract embedded text.]")]


def parse_audio(path: str) -> List[Tuple[int, str]]:
    """
    Speech-to-text stub. Wire up faster-whisper / OpenAI Whisper in
    production (see backend/utils/parsers.py). Kept out of the default
    MVP install to keep the container light for the hackathon demo.
    """
    return [(1, f"[Audio document: {os.path.basename(path)} — transcription requires a Whisper "
                 f"worker. Configure WHISPER_MODEL in .env to enable.]")]


EXT_MAP = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".txt": parse_txt,
    ".csv": parse_csv,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".png": parse_image,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".mp3": parse_audio,
    ".wav": parse_audio,
    ".m4a": parse_audio,
}


def detect_file_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    mapping = {
        ".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".csv": "csv",
        ".xlsx": "xlsx", ".xls": "xlsx", ".png": "image", ".jpg": "image",
        ".jpeg": "image", ".mp3": "audio", ".wav": "audio", ".m4a": "audio",
    }
    return mapping.get(ext, "unknown")


def extract_text(path: str) -> List[Tuple[int, str]]:
    ext = os.path.splitext(path)[1].lower()
    parser = EXT_MAP.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(path)


def clean_text(text: str) -> str:
    lines = [l.strip() for l in text.splitlines()]
    lines = [l for l in lines if l]
    return "\n".join(lines)


def chunk_text(pages: List[Tuple[int, str]], chunk_size: int = 900, overlap: int = 150) -> List[dict]:
    """Sliding-window chunker that preserves page numbers for citations."""
    chunks = []
    idx = 0
    for page_num, raw in pages:
        text = clean_text(raw)
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            chunks.append({"text": chunk, "page": page_num, "index": idx})
            idx += 1
            if end == len(text):
                break
            start = end - overlap
    return chunks
